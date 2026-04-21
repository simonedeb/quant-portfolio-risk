from docx import Document
from docx.shared import Pt, RGBColor

doc = Document(r'c:\Users\simon\OneDrive\Desktop\MSC FINANC\MSc_Finance_CV_v2.docx')

GRAY = RGBColor(0x55, 0x55, 0x55)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(0)
r = p.add_run(
    'I hereby authorise the processing of my personal data included in this CV for recruitment and selection purposes, '
    'in accordance with the EU General Data Protection Regulation (GDPR) 2016/679.'
)
r.font.size = Pt(8)
r.font.color.rgb = GRAY
r.italic = True

doc.save(r'c:\Users\simon\OneDrive\Desktop\MSC FINANC\MSc_Finance_CV_v2.docx')
print('Done.')
