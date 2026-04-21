from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document(r'c:\Users\simon\OneDrive\Desktop\MSC FINANC\MSc_Finance_CV_v3.docx')

# Balanced margins — tight but readable
for section in doc.sections:
    section.top_margin    = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin   = Cm(1.6)
    section.right_margin  = Cm(1.6)

for para in doc.paragraphs:
    pf = para.paragraph_format
    # Restore small breathing room
    if pf.space_before is not None and pf.space_before.pt == 0:
        pf.space_before = Pt(2)
    pf.space_after = Pt(1)
    pf.line_spacing = Pt(12)

    for run in para.runs:
        size = run.font.size
        if size:
            pts = size.pt
            if pts >= 20:
                run.font.size = Pt(20)
            elif pts >= 11:
                run.font.size = Pt(10.5)
            elif pts >= 10:
                run.font.size = Pt(10)
            elif pts >= 9:
                run.font.size = Pt(9.5)
            else:
                run.font.size = Pt(8)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(3)
                para.paragraph_format.space_after  = Pt(1)
                para.paragraph_format.line_spacing = Pt(12)
                for run in para.runs:
                    size = run.font.size
                    if size:
                        pts = size.pt
                        if pts >= 20:
                            run.font.size = Pt(20)
                        elif pts >= 11:
                            run.font.size = Pt(10.5)
                        elif pts >= 10:
                            run.font.size = Pt(10)
                        elif pts >= 9:
                            run.font.size = Pt(9.5)
                        else:
                            run.font.size = Pt(8)

doc.save(r'c:\Users\simon\OneDrive\Desktop\MSC FINANC\MSc_Finance_CV_v3.docx')
print('Done.')
