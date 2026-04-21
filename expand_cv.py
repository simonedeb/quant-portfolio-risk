from docx import Document
from docx.shared import Pt, Cm

doc = Document(r'c:\Users\simon\OneDrive\Desktop\MSC FINANC\MSc_Finance_CV_v2.docx')

# Slightly looser margins
for section in doc.sections:
    section.top_margin    = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin   = Cm(1.8)
    section.right_margin  = Cm(1.8)

for para in doc.paragraphs:
    pf = para.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after  = Pt(2)
    pf.line_spacing = Pt(12.5)

    for run in para.runs:
        size = run.font.size
        if size:
            pts = size.pt
            if pts >= 18:
                run.font.size = Pt(22)
            elif pts >= 10:
                run.font.size = Pt(10.5)
            elif pts >= 9:
                run.font.size = Pt(10)
            else:
                run.font.size = Pt(8.5)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(3)
                para.paragraph_format.space_after  = Pt(2)
                para.paragraph_format.line_spacing = Pt(12.5)
                for run in para.runs:
                    size = run.font.size
                    if size:
                        pts = size.pt
                        if pts >= 18:
                            run.font.size = Pt(22)
                        elif pts >= 10:
                            run.font.size = Pt(10.5)
                        elif pts >= 9:
                            run.font.size = Pt(10)
                        else:
                            run.font.size = Pt(8.5)

doc.save(r'c:\Users\simon\OneDrive\Desktop\MSC FINANC\MSc_Finance_CV_v2.docx')
print('Done.')
