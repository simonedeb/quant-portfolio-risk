from docx import Document
import re

doc = Document(r'c:\Users\simon\OneDrive\Desktop\MSC FINANC\MSc_Finance_CV_v2.docx')

for para in doc.paragraphs:
    for run in para.runs:
        t = run.text

        # Remove trailing separator after last contact item
        if t.strip().endswith('|') and 'Italy' in t:
            run.text = t.rstrip().rstrip('|').rstrip()

        # Remove square brackets around university/school names
        if re.search(r'\[.*?\]', t):
            run.text = re.sub(r'\[([^\]]+)\]', r'\1', t)

    # Fix paragraph with lone "|"
    full_text = ''.join(r.text for r in para.runs)
    if full_text.strip() == '|':
        for run in para.runs:
            run.text = ''

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    t = run.text
                    if re.search(r'\[.*?\]', t):
                        run.text = re.sub(r'\[([^\]]+)\]', r'\1', t)

doc.save(r'c:\Users\simon\OneDrive\Desktop\MSC FINANC\MSc_Finance_CV_v2.docx')
print('Fixed.')
