from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Tight margins for 1 page
for section in doc.sections:
    section.top_margin    = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin   = Cm(1.6)
    section.right_margin  = Cm(1.6)

DARK   = RGBColor(0x1A, 0x1A, 0x2E)
ACCENT = RGBColor(0x16, 0x21, 0x3E)
GRAY   = RGBColor(0x55, 0x55, 0x55)

def remove_table_borders(table):
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBorders.append(el)
    tblPr.append(tblBorders)

def add_section_header(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = ACCENT
    run.font.name = 'Calibri'
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '162136')
    pBdr.append(bot)
    pPr.append(pBdr)

def add_entry(doc, left_bold, left_normal, right_text, sub_left=None, sub_right=None, bullets=None):
    t = doc.add_table(rows=1, cols=2)
    remove_table_borders(t)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.columns[0].width = Inches(4.6)
    t.columns[1].width = Inches(1.7)

    p_l = t.cell(0,0).paragraphs[0]
    p_l.paragraph_format.space_before = Pt(4)
    p_l.paragraph_format.space_after  = Pt(0)
    r1 = p_l.add_run(left_bold)
    r1.bold = True
    r1.font.size = Pt(9.5)
    r1.font.color.rgb = DARK
    if left_normal:
        r2 = p_l.add_run(f'  {left_normal}')
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = GRAY

    p_r = t.cell(0,1).paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_r.paragraph_format.space_before = Pt(4)
    p_r.paragraph_format.space_after  = Pt(0)
    r3 = p_r.add_run(right_text)
    r3.font.size = Pt(9)
    r3.font.color.rgb = GRAY
    r3.italic = True

    if sub_left or sub_right:
        t2 = doc.add_table(rows=1, cols=2)
        remove_table_borders(t2)
        t2.columns[0].width = Inches(4.6)
        t2.columns[1].width = Inches(1.7)
        sl = t2.cell(0,0).paragraphs[0]
        sl.paragraph_format.space_before = Pt(0)
        sl.paragraph_format.space_after  = Pt(0)
        r_sl = sl.add_run(sub_left or '')
        r_sl.italic = True
        r_sl.font.size = Pt(9)
        r_sl.font.color.rgb = GRAY
        if sub_right:
            sr = t2.cell(0,1).paragraphs[0]
            sr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r_sr = sr.add_run(sub_right)
            r_sr.font.size = Pt(9)
            r_sr.font.color.rgb = GRAY

    if bullets:
        for b in bullets:
            p_b = doc.add_paragraph(style='List Bullet')
            p_b.paragraph_format.space_before = Pt(0)
            p_b.paragraph_format.space_after  = Pt(0)
            p_b.paragraph_format.left_indent  = Cm(0.4)
            rb = p_b.add_run(b)
            rb.font.size = Pt(9)
            rb.font.color.rgb = GRAY

# ─── HEADER ───────────────────────────────────────────────────────────────────
p_name = doc.add_paragraph()
p_name.paragraph_format.space_before = Pt(0)
p_name.paragraph_format.space_after  = Pt(1)
r_name = p_name.add_run('FIRST NAME  LAST NAME')
r_name.bold = True
r_name.font.size = Pt(22)
r_name.font.color.rgb = DARK
r_name.font.name = 'Calibri'

p_tag = doc.add_paragraph()
p_tag.paragraph_format.space_before = Pt(0)
p_tag.paragraph_format.space_after  = Pt(3)
r_tag = p_tag.add_run('BSc Business Administration — Major in Banking & Finance  |  MSc Finance Candidate 2025')
r_tag.font.size = Pt(9.5)
r_tag.font.color.rgb = ACCENT
r_tag.italic = True

p_contact = doc.add_paragraph()
p_contact.paragraph_format.space_before = Pt(0)
p_contact.paragraph_format.space_after  = Pt(0)
contact_items = ['email@email.com', '+39 000 000 0000', 'linkedin.com/in/yourprofile', 'City, Country']
for i, item in enumerate(contact_items):
    sep = '   |   ' if i < len(contact_items) - 1 else ''
    r = p_contact.add_run(f'{item}{sep}')
    r.font.size = Pt(9)
    r.font.color.rgb = GRAY

# Divider
p_div = doc.add_paragraph()
p_div.paragraph_format.space_before = Pt(3)
p_div.paragraph_format.space_after  = Pt(2)
pPr  = p_div._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bot  = OxmlElement('w:bottom')
bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'12')
bot.set(qn('w:space'),'1');    bot.set(qn('w:color'),'162136')
pBdr.append(bot); pPr.append(pBdr)

# ─── EDUCATION ────────────────────────────────────────────────────────────────
add_section_header(doc, 'Education')

add_entry(doc,
    'Bachelor of Science in Business Administration',
    '— Major in Banking & Finance',
    'Sep 20XX – Jul/Sep 2025 (Expected)',
    sub_left='[University Name] — [City, Country]',
    sub_right='Major GPA: 29.6/30  |  Overall: 26.08/30',
    bullets=[
        'Coursework: Statistics [30/30], Mathematical Finance [30/30], Financial Accounting [30L/30], Technical Analysis & Stock Trading [30L/30]',
    ]
)

add_entry(doc,
    'Scientific High School Diploma',
    '',
    'Jun 20XX',
    sub_left='[High School Name] — [City, Country]',
    sub_right='Final Grade: 100 / 100',
    bullets=[
        'Led peer study groups and coordinated class initiatives; developed leadership, communication, and organisational skills over a 5-year programme',
    ]
)

# ─── INDEPENDENT RESEARCH & PROJECTS ─────────────────────────────────────────
add_section_header(doc, 'Independent Research & Projects')

add_entry(doc,
    'Equity Valuation — Ferrari N.V. (RACE.MI)',
    '',
    '2025 – Present',
    sub_left='Independent Project  ·  Discounted Cash Flow Analysis',
    bullets=[
        'Built multi-stage DCF in Excel/Python; estimated WACC via CAPM; performed sensitivity analysis on terminal growth rate and discount rate; stress-tested under bull/base/bear scenarios using Ferrari annual reports and sell-side consensus estimates',
    ]
)

add_entry(doc,
    'Portfolio Optimisation — Large-Cap European Equities',
    '',
    '2025 – Present',
    sub_left='Independent Project  ·  Value-at-Risk & Mean-Variance Optimisation',
    bullets=[
        'Constructed diversified 10-stock EU portfolio in Python; derived efficient frontier via Mean-Variance Optimisation; implemented Historical and Parametric VaR at 95%/99% confidence; back-tested estimates against realised returns',
    ]
)

# ─── EXTRACURRICULAR & LEADERSHIP ─────────────────────────────────────────────
add_section_header(doc, 'Extracurricular Activities & Leadership')

add_entry(doc,
    'Intraday Trading Competition',
    '',
    'Summer 20XX',
    sub_left='[Programme Name] — Summer Intensive Trading Week  ·  [City, Country]',
    bullets=[
        'First direct exposure to live financial markets in a competitive setting; executed intraday strategies on high-volatility assets across a 5-day simulation',
        'Applied stop-loss discipline, position sizing, and risk-reward frameworks under strict capital constraints and daily loss limits',
    ]
)

# ─── SKILLS & CERTIFICATIONS ──────────────────────────────────────────────────
add_section_header(doc, 'Skills & Certifications')

skills_table = doc.add_table(rows=4, cols=2)
remove_table_borders(skills_table)
skills_table.columns[0].width = Inches(1.2)
skills_table.columns[1].width = Inches(5.1)

rows_data = [
    ('Technical',       'Python  |  Excel (Advanced)  |  Microsoft Office Suite  |  Bloomberg Terminal'),
    ('Languages',       'Italian: Native  |  English: Fluent — IELTS 7.0 (C1)'),
    ('Certifications',  'Bloomberg Market Concepts (BMC) — Certified'),
    ('Sports',          'Competitive football (2016–2023)  |  Padel & Tennis (Amateur)  |  Kickboxing  |  Fitness'),
]

for i, (label, value) in enumerate(rows_data):
    row = skills_table.rows[i]
    cl = row.cells[0].paragraphs[0]
    cl.paragraph_format.space_before = Pt(2)
    rl = cl.add_run(label)
    rl.bold = True
    rl.font.size = Pt(9)
    rl.font.color.rgb = DARK
    cv = row.cells[1].paragraphs[0]
    cv.paragraph_format.space_before = Pt(2)
    rv = cv.add_run(value)
    rv.font.size = Pt(9)
    rv.font.color.rgb = GRAY

# ─── GDPR CONSENT ─────────────────────────────────────────────────────────────
p_gdpr = doc.add_paragraph()
p_gdpr.paragraph_format.space_before = Pt(6)
p_gdpr.paragraph_format.space_after  = Pt(0)
r_gdpr = p_gdpr.add_run(
    'I hereby authorise the processing of my personal data included in this CV for recruitment and selection purposes, '
    'in accordance with the EU General Data Protection Regulation (GDPR) 2016/679.'
)
r_gdpr.font.size = Pt(8)
r_gdpr.font.color.rgb = GRAY
r_gdpr.italic = True

# Font cleanup
for para in doc.paragraphs:
    for run in para.runs:
        if not run.font.name:
            run.font.name = 'Calibri'

output_path = r'c:\Users\simon\OneDrive\Desktop\MSC FINANC\MSc_Finance_CV_v3.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
