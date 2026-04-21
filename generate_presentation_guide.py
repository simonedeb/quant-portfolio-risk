from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Colours
NAVY   = RGBColor(0x0D, 0x1B, 0x3E)
ACCENT = RGBColor(0x00, 0x6A, 0xC1)
GRAY   = RGBColor(0x44, 0x44, 0x44)
RED    = RGBColor(0xC0, 0x00, 0x00)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLACK  = RGBColor(0x00, 0x00, 0x00)

doc = Document()

# Margins
for s in doc.sections:
    s.top_margin = s.bottom_margin = Cm(2.0)
    s.left_margin = s.right_margin = Cm(2.5)

def para(text='', size=10, bold=False, italic=False,
         color=BLACK, space_before=4, space_after=3,
         indent=0, font='Calibri'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if text:
        r = p.add_run(text)
        r.bold = bold; r.italic = italic
        r.font.size = Pt(size)
        r.font.color.rgb = color
        r.font.name = font
    return p

def h1(text):
    para(text, size=18, bold=True, color=NAVY, space_before=0, space_after=4)

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(2)
    # Horizontal rule above
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top  = OxmlElement('w:top')
    top.set(qn('w:val'), 'single'); top.set(qn('w:sz'), '4')
    top.set(qn('w:space'), '1');    top.set(qn('w:color'), '006AC1')
    pBdr.append(top); pPr.append(pBdr)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(13)
    r.font.color.rgb = ACCENT; r.font.name = 'Calibri'

def quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.right_indent = Cm(1.0)
    # Left border
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    lft  = OxmlElement('w:left')
    lft.set(qn('w:val'), 'single'); lft.set(qn('w:sz'), '12')
    lft.set(qn('w:space'), '8');    lft.set(qn('w:color'), '006AC1')
    pBdr.append(lft); pPr.append(pBdr)
    r = p.add_run(f'"{text}"')
    r.italic = True; r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x22, 0x22, 0x66); r.font.name = 'Calibri'

def bullet(text, indent=0.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(indent)
    r = p.add_run(text)
    r.font.size = Pt(10); r.font.color.rgb = GRAY; r.font.name = 'Calibri'

def label(p, text, bold=True, color=NAVY):
    r = p.add_run(text)
    r.bold = bold; r.font.size = Pt(10)
    r.font.color.rgb = color; r.font.name = 'Calibri'

def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        run = hdr[i].paragraphs[0].runs[0]
        run.bold = True; run.font.size = Pt(9.5)
        run.font.color.rgb = WHITE; run.font.name = 'Calibri'
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '0D1B3E')
        tcPr.append(shd)
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        fill = 'F0F4FA' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            cells[ci].text = val
            run = cells[ci].paragraphs[0].runs[0]
            run.font.size = Pt(9.5); run.font.name = 'Calibri'
            run.font.color.rgb = GRAY
            tc = cells[ci]._tc
            tcPr = tc.get_or_add_tcPr()
            shd  = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            tcPr.append(shd)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ════════════════════════════════════════════════════════════════════════════
# TITLE
# ════════════════════════════════════════════════════════════════════════════
h1('Quantitative Risk & Portfolio Management')
para('Independent Research Project  —  MSc Finance Applications',
     size=11, color=GRAY, space_before=0, space_after=2)
para('Presentation Guide  |  Simon De Bartolo',
     size=10, italic=True, color=GRAY, space_before=0, space_after=10)

# ════════════════════════════════════════════════════════════════════════════
# OVERVIEW
# ════════════════════════════════════════════════════════════════════════════
h2('Project Overview — Opening Statement')
quote('I built an end-to-end quantitative risk and portfolio management system in Python, '
      'covering the full workflow a risk analyst would follow: data collection, return analysis, '
      'risk measurement, model validation, portfolio optimisation, and stress testing. '
      'The project uses 10 large-cap European equities over a 3-year period.')

# ════════════════════════════════════════════════════════════════════════════
# STEP 1
# ════════════════════════════════════════════════════════════════════════════
h2('Step 1 — Data Collection')
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
label(p, 'What: '); label(p, 'Downloaded daily adjusted closing prices for 10 EU stocks (ASML, LVMH, SAP, Siemens, TotalEnergies…) from Yahoo Finance via ', bold=False, color=GRAY); label(p, 'yfinance', color=ACCENT); label(p, ', 2022–2025.', bold=False, color=GRAY)
p2 = doc.add_paragraph(); p2.paragraph_format.space_before = Pt(2)
label(p2, 'Why: '); label(p2, 'Adjusted prices account for dividends and splits — essential for accurate return calculation.', bold=False, color=GRAY)
para('Say to committee:', bold=True, size=9.5, color=NAVY, space_before=6, space_after=2)
quote('I used Python\'s yfinance library to pull 3 years of daily price data for 10 diversified European large-caps across sectors: tech, luxury, energy, financials. Using adjusted prices ensures the return series is not distorted by corporate actions.')

# ════════════════════════════════════════════════════════════════════════════
# STEP 2
# ════════════════════════════════════════════════════════════════════════════
h2('Step 2 — Log Returns & Descriptive Statistics')
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
label(p, 'Formula: '); label(p, 'r_t = ln(P_t / P_{t-1})     Annualised mean x252     Annualised volatility x sqrt(252)', bold=False, color=GRAY)
p2 = doc.add_paragraph(); p2.paragraph_format.space_before = Pt(2)
label(p2, 'Why log returns: '); label(p2, 'Additive across time, approximately normally distributed, standard in quantitative finance.', bold=False, color=GRAY)
p3 = doc.add_paragraph(); p3.paragraph_format.space_before = Pt(2)
label(p3, 'Key output: '); label(p3, 'Correlation matrix + cumulative return chart showing cross-asset diversification.', bold=False, color=GRAY)
para('Say to committee:', bold=True, size=9.5, color=NAVY, space_before=6, space_after=2)
quote('Log returns are the industry standard because they are time-additive and better approximate a normal distribution than simple returns. I annualised by multiplying mean daily return by 252 trading days and volatility by the square root of 252. The correlation matrix showed that most pairs had moderate positive correlation — confirming diversification benefit but also systemic exposure.')

# ════════════════════════════════════════════════════════════════════════════
# STEP 3
# ════════════════════════════════════════════════════════════════════════════
h2('Step 3 — Value at Risk (VaR)')
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
label(p, 'Definition: '); label(p, 'Maximum expected loss over one day at a given confidence level.', bold=False, color=GRAY)
para('Three methods compared:', bold=True, size=10, color=NAVY, space_before=6, space_after=2)
add_table(
    ['Method', 'How', 'Assumption'],
    [
        ['Historical', '5th percentile of actual daily returns', 'None — non-parametric'],
        ['Parametric', 'z-score x sigma using normal distribution', 'Returns ~ Normal'],
        ['Monte Carlo', '10,000 simulated normal draws', 'Returns ~ Normal'],
    ]
)
para('Say to committee:', bold=True, size=9.5, color=NAVY, space_before=6, space_after=2)
quote('VaR answers: what is the maximum loss I should expect on 95% of trading days? Historical VaR makes no distributional assumption — it reads directly from empirical data. Parametric VaR is faster but assumes normality, which financial returns often violate due to fat tails. Monte Carlo confirms the parametric estimate with simulation. At 95%, historical VaR was around -0.9%, meaning on 1 in 20 days we expect to lose at least that amount.')

# ════════════════════════════════════════════════════════════════════════════
# STEP 4
# ════════════════════════════════════════════════════════════════════════════
h2('Step 4 — Conditional VaR (Expected Shortfall / CVaR)')
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
label(p, 'Definition: '); label(p, 'Average return in the worst 5% of days — the mean of the tail beyond VaR.', bold=False, color=GRAY)
p2 = doc.add_paragraph(); p2.paragraph_format.space_before = Pt(2)
label(p2, 'Regulatory context: '); label(p2, 'CVaR replaced VaR as primary capital metric under Basel III / FRTB (2019).', bold=False, color=RED)
p3 = doc.add_paragraph(); p3.paragraph_format.space_before = Pt(2)
label(p3, 'Result: '); label(p3, 'CVaR 95% ~ -1.4% vs VaR 95% ~ -0.9% — tail is worse than the threshold suggests.', bold=False, color=GRAY)
para('Say to committee:', bold=True, size=9.5, color=NAVY, space_before=6, space_after=2)
quote('VaR has a well-known limitation: it says nothing about how bad losses get once you cross the threshold. CVaR, or Expected Shortfall, fills this gap by averaging losses in the worst tail. Regulators recognised this — Basel III\'s Fundamental Review of the Trading Book replaced VaR with CVaR as the primary capital charge metric. In my portfolio, CVaR at 95% was approximately -1.4%, meaningfully worse than -0.9% VaR, confirming non-trivial tail risk.')

# ════════════════════════════════════════════════════════════════════════════
# STEP 5
# ════════════════════════════════════════════════════════════════════════════
h2('Step 5 — Backtesting (Kupiec Test)')
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
label(p, 'What: '); label(p, 'Counted days where actual loss exceeded VaR estimate (breaches). Applied Kupiec POF test.', bold=False, color=GRAY)
p2 = doc.add_paragraph(); p2.paragraph_format.space_before = Pt(2)
label(p2, 'Kupiec test: '); label(p2, 'Likelihood-ratio test. H0: observed breach rate = expected rate (5% at 95% confidence).', bold=False, color=GRAY)
para('Interpretation:', bold=True, size=10, color=NAVY, space_before=6, space_after=2)
bullet('Breaches >> 5%  →  model underestimates risk')
bullet('Breaches << 5%  →  model is too conservative')
bullet('p-value > 0.05  →  fail to reject H0  →  model is well-calibrated')
para('Say to committee:', bold=True, size=9.5, color=NAVY, space_before=6, space_after=2)
quote('A VaR model is only useful if it is accurate. Backtesting counts actual exceedances and compares them to the expected frequency. I used the Kupiec test, a formal likelihood-ratio hypothesis test. A p-value above 0.05 means I cannot reject the null that the model is correctly calibrated. This is exactly the validation approach regulators require banks to perform daily on their trading books.')

# ════════════════════════════════════════════════════════════════════════════
# STEP 6
# ════════════════════════════════════════════════════════════════════════════
h2('Step 6 — Portfolio Optimisation (Markowitz / Mean-Variance)')
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
label(p, 'Framework: '); label(p, 'Markowitz (1952) Mean-Variance Optimisation. Minimise portfolio variance subject to return and weight constraints. Solver: SLSQP (scipy.optimize).', bold=False, color=GRAY)
para('Four portfolios constructed:', bold=True, size=10, color=NAVY, space_before=6, space_after=2)
add_table(
    ['Portfolio', 'Objective'],
    [
        ['Equal Weight',  'Baseline — 1/n allocation across all assets'],
        ['Min Variance',  'Minimise sigma — lowest possible risk'],
        ['Max Sharpe',    'Maximise return / risk ratio (Sharpe ratio)'],
        ['Min VaR',       'Minimise historical 5th-percentile tail loss'],
    ]
)
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
label(p, 'Efficient frontier: '); label(p, '100 optimised portfolios across target return range, plus 5,000 Monte Carlo random portfolios (Dirichlet weights) to visualise feasible set.', bold=False, color=GRAY)
p2 = doc.add_paragraph(); p2.paragraph_format.space_before = Pt(2)
label(p2, 'Capital Market Line: '); label(p2, 'Tangent line from risk-free rate (3.5%) to Max Sharpe point — the theoretically optimal combination of risky assets.', bold=False, color=GRAY)
para('Say to committee:', bold=True, size=9.5, color=NAVY, space_before=6, space_after=2)
quote('Markowitz showed that investors should not evaluate assets in isolation but consider how they co-move. By solving a constrained quadratic optimisation, I traced the efficient frontier — the set of portfolios offering maximum return for each level of risk. The Max Sharpe portfolio is the theoretically optimal risky portfolio; combined with a risk-free asset it forms the Capital Market Line. The 5,000 Monte Carlo simulations visualise the entire feasible investment space — the frontier sits at the upper-left boundary of the cloud.')

# ════════════════════════════════════════════════════════════════════════════
# STEP 7
# ════════════════════════════════════════════════════════════════════════════
h2('Step 7 — Stress Testing')
para('Three scenarios tested across all four portfolios:', bold=True, size=10, color=NAVY, space_before=4, space_after=2)
add_table(
    ['Scenario', 'Source', 'Reference Loss'],
    [
        ['2022 Rate-Hike Selloff', 'Actual return data  Jan-Oct 2022', 'Real data'],
        ['COVID-like shock',       'Simulated tail draw, scaled to EW reference', '-35%'],
        ['GFC-like shock',         'Simulated tail draw, scaled to EW reference', '-50%'],
    ]
)
para('Key finding:', bold=True, size=10, color=NAVY, space_before=6, space_after=2)
bullet('Min Variance consistently outperforms under stress in every scenario')
bullet('2022: Min Variance -5.1%  vs  Equal Weight -9.5%  vs  Max Sharpe -2.4% (but MDD -20.7%)')
bullet('COVID-like: Min Variance -17.4%  vs  Max Sharpe -31.7%')
bullet('Max Sharpe concentrates in few high-return assets — high upside, severe drawdown in crashes')
para('Say to committee:', bold=True, size=9.5, color=NAVY, space_before=6, space_after=2)
quote('VaR and CVaR measure normal market conditions. Stress testing asks: what happens in a genuine crisis? The 2022 selloff was driven by the fastest rate-hiking cycle in decades — it is real data in my sample. For COVID and the 2008 GFC, which predate my data, I constructed hypothetical scenarios by sampling from the empirical tail and scaling to reference losses of 35% and 50%. The results clearly show that the Min Variance portfolio — concentrated in low-volatility, low-correlation assets — provides meaningful downside protection precisely when it is most needed.')

# ════════════════════════════════════════════════════════════════════════════
# CLOSING
# ════════════════════════════════════════════════════════════════════════════
h2('Closing Statement')
quote('This project demonstrates the full quantitative risk management pipeline: from raw market data to model validation and scenario analysis. The tools — Python, NumPy, SciPy, pandas — are standard in industry quant roles. More importantly, each step is grounded in theory: Markowitz optimisation, Basel III risk metrics, and statistical hypothesis testing. My goal was not just to run code but to understand what each number means and why it matters for portfolio decisions.')

# ════════════════════════════════════════════════════════════════════════════
# KEY CONCEPTS GLOSSARY
# ════════════════════════════════════════════════════════════════════════════
h2('Key Concepts — Quick Reference (memorise these)')
add_table(
    ['Concept', 'One-sentence definition'],
    [
        ['Log Return',           'r = ln(P_t / P_{t-1}) — time-additive, approximately normal'],
        ['VaR (Value at Risk)',   'Maximum loss not exceeded on X% of days at a given confidence level'],
        ['CVaR / Expected Shortfall', 'Average loss in the worst tail beyond VaR — Basel III standard'],
        ['Kupiec Test',          'Likelihood-ratio test validating that VaR breach frequency matches model expectation'],
        ['Efficient Frontier',   'Set of portfolios maximising return for each level of risk (Markowitz)'],
        ['Sharpe Ratio',         '(Return - Risk-free rate) / Volatility — risk-adjusted performance measure'],
        ['Capital Market Line',  'Tangent from risk-free rate to Max Sharpe portfolio — optimal asset allocation line'],
        ['Max Drawdown',         'Largest peak-to-trough loss over a period — measures worst-case sustained loss'],
        ['Stress Testing',       'Portfolio evaluation under extreme but plausible market scenarios beyond normal VaR'],
        ['SLSQP',                'Sequential Least Squares Programming — constrained numerical optimisation algorithm'],
    ]
)

path = r'c:\Users\simon\OneDrive\Desktop\MSC FINANC\Presentation_Guide.docx'
doc.save(path)
print(f'Saved: {path}')
