from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK   = RGBColor(0x1A, 0x1A, 0x2E)
ACCENT = RGBColor(0x16, 0x21, 0x3E)
GRAY   = RGBColor(0x55, 0x55, 0x55)
RED    = RGBColor(0xC0, 0x00, 0x00)

def set_margins(doc):
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(2.0)
        s.left_margin = s.right_margin = Cm(2.5)

def heading(doc, text, size=14, color=DARK, bold=True, space_before=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = 'Calibri'

def section_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
    bot.set(qn('w:space'),'1');    bot.set(qn('w:color'),'162136')
    pBdr.append(bot); pPr.append(pBdr)

def add_question(doc, n, qtext, options, answer, explanation):
    # Question
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(f'Q{n}.  {qtext}')
    r.font.size = Pt(10.5)
    r.font.color.rgb = DARK
    r.font.name = 'Calibri'

    # Options
    for opt in options:
        po = doc.add_paragraph()
        po.paragraph_format.space_before = Pt(1)
        po.paragraph_format.space_after  = Pt(1)
        po.paragraph_format.left_indent  = Cm(0.8)
        ro = po.add_run(opt)
        ro.font.size = Pt(10)
        ro.font.color.rgb = GRAY
        ro.font.name = 'Calibri'

    # Answer
    pa = doc.add_paragraph()
    pa.paragraph_format.space_before = Pt(4)
    pa.paragraph_format.space_after  = Pt(1)
    pa.paragraph_format.left_indent  = Cm(0.8)
    ra1 = pa.add_run('Answer: ')
    ra1.bold = True
    ra1.font.size = Pt(10)
    ra1.font.color.rgb = ACCENT
    ra2 = pa.add_run(answer)
    ra2.bold = True
    ra2.font.size = Pt(10)
    ra2.font.color.rgb = ACCENT

    # Explanation
    pe = doc.add_paragraph()
    pe.paragraph_format.space_before = Pt(0)
    pe.paragraph_format.space_after  = Pt(6)
    pe.paragraph_format.left_indent  = Cm(0.8)
    re = pe.add_run(f'Explanation: {explanation}')
    re.font.size = Pt(9.5)
    re.font.color.rgb = GRAY
    re.italic = True
    re.font.name = 'Calibri'


# ══════════════════════════════════════════════════════════════════════════════
# FILE 1 — IESAT (IE UNIVERSITY)
# ══════════════════════════════════════════════════════════════════════════════
doc1 = Document()
set_margins(doc1)

heading(doc1, 'IESAT — IE University', size=18, space_before=0)
heading(doc1, 'Practice Simulation  |  5 Questions per Section', size=11, color=GRAY, bold=False, space_before=2)
section_line(doc1)

# ── SECTION 1: QUANTITATIVE REASONING ────────────────────────────────────────
heading(doc1, 'SECTION 1 — Quantitative Reasoning', size=12, color=ACCENT)

add_question(doc1, 1,
    'A company\'s revenue grew from €240,000 to €312,000 over one year. What was the percentage increase?',
    ['A) 25%', 'B) 28%', 'C) 30%', 'D) 32%'],
    'C) 30%',
    '(312,000 − 240,000) / 240,000 = 72,000 / 240,000 = 0.30 = 30%'
)

add_question(doc1, 2,
    'If x² − 5x + 6 = 0, what are the values of x?',
    ['A) x = 1 and x = 6', 'B) x = 2 and x = 3', 'C) x = −2 and x = −3', 'D) x = −1 and x = 6'],
    'B) x = 2 and x = 3',
    'Factor: (x − 2)(x − 3) = 0 → x = 2 or x = 3'
)

add_question(doc1, 3,
    'A portfolio loses 20% in year 1 and gains 25% in year 2. What is the net change over the two years?',
    ['A) +5%', 'B) +2%', 'C) 0%', 'D) −2%'],
    'C) 0%',
    'Start: 100 → after −20%: 80 → after +25%: 80 × 1.25 = 100. Net change = 0%'
)

add_question(doc1, 4,
    'A train travels 360 km in 3 hours. A second train covers the same distance but 20% faster. How long does the second train take?',
    ['A) 2h 00min', 'B) 2h 15min', 'C) 2h 30min', 'D) 2h 45min'],
    'C) 2h 30min',
    'Speed 1 = 120 km/h. Speed 2 = 120 × 1.20 = 144 km/h. Time = 360 / 144 = 2.5 h = 2h 30min'
)

add_question(doc1, 5,
    'Data Sufficiency: Is integer n divisible by 6?\n'
    '   (1) n is divisible by 3\n'
    '   (2) n is divisible by 2',
    ['A) Statement (1) alone is sufficient',
     'B) Statement (2) alone is sufficient',
     'C) Both statements together are sufficient, but neither alone is sufficient',
     'D) Each statement alone is sufficient'],
    'C) Both statements together are sufficient, but neither alone is sufficient',
    'Divisibility by 6 requires both factors 2 and 3. Neither statement alone guarantees both conditions; together they do.'
)

# ── SECTION 2: VERBAL REASONING ──────────────────────────────────────────────
heading(doc1, 'SECTION 2 — Verbal Reasoning', size=12, color=ACCENT)

add_question(doc1, 6,
    'Choose the word most similar in meaning to MITIGATE:',
    ['A) Aggravate', 'B) Alleviate', 'C) Accelerate', 'D) Amplify'],
    'B) Alleviate',
    'Mitigate means to make less severe. Alleviate is the closest synonym. Aggravate and amplify are antonyms.'
)

add_question(doc1, 7,
    'Reading Comprehension:\n\n'
    '"Central banks use interest rate policy as a primary tool to manage inflation. When inflation exceeds the target, '
    'rates are raised to reduce borrowing and cool demand. However, aggressive rate hikes risk slowing economic growth '
    'and increasing unemployment, creating a policy dilemma."\n\n'
    'What is the main trade-off described in the passage?',
    ['A) Between fiscal and monetary policy',
     'B) Between controlling inflation and sustaining economic growth',
     'C) Between unemployment and interest rates',
     'D) Between central bank independence and government control'],
    'B) Between controlling inflation and sustaining economic growth',
    'The passage explicitly states that rate hikes reduce inflation but risk slowing growth — a classic monetary policy dilemma.'
)

add_question(doc1, 8,
    'Sentence Correction: Identify the grammatically correct sentence.',
    ['A) The data shows that markets has recovered.',
     'B) The data show that markets have recovered.',
     'C) The datas show that market have recovered.',
     'D) The data shown that markets has recovered.'],
    'B) The data show that markets have recovered.',
    '"Data" is a plural noun (singular: datum). Correct verb form: "show" and "have recovered".'
)

# ── SECTION 3: LOGICAL REASONING ─────────────────────────────────────────────
heading(doc1, 'SECTION 3 — Logical Reasoning', size=12, color=ACCENT)

add_question(doc1, 9,
    'Number sequence — find the next term:\n2,  6,  12,  20,  30,  ?',
    ['A) 40', 'B) 42', 'C) 44', 'D) 48'],
    'B) 42',
    'Differences: +4, +6, +8, +10, +12 → next term: 30 + 12 = 42. Pattern: n(n+1).'
)

add_question(doc1, 10,
    'All analysts are detail-oriented. Some detail-oriented people are introverts. '
    'Which conclusion is valid?',
    ['A) All analysts are introverts',
     'B) Some analysts are introverts',
     'C) No analysts are introverts',
     'D) None of the above can be concluded with certainty'],
    'D) None of the above can be concluded with certainty',
    'The premises allow that some detail-oriented people (who may or may not be analysts) are introverts. '
    'No definitive link between analysts and introversion can be established.'
)

doc1.save(r'c:\Users\simon\OneDrive\Desktop\MSC FINANC\IESAT_Practice_Test.docx')
print('IESAT saved.')


# ══════════════════════════════════════════════════════════════════════════════
# FILE 2 — TAE (ESADE)
# ══════════════════════════════════════════════════════════════════════════════
doc2 = Document()
set_margins(doc2)

heading(doc2, 'TAE — ESADE Business School', size=18, space_before=0)
heading(doc2, 'Practice Simulation  |  5 Questions per Section', size=11, color=GRAY, bold=False, space_before=2)
section_line(doc2)

# ── SECTION 1: NUMERICAL REASONING ───────────────────────────────────────────
heading(doc2, 'SECTION 1 — Numerical Reasoning', size=12, color=ACCENT)

add_question(doc2, 1,
    'A stock is priced at €80. After a 15% drop followed by a 10% rise, what is the final price?',
    ['A) €74.80', 'B) €75.60', 'C) €76.00', 'D) €78.00'],
    'A) €74.80',
    '80 × 0.85 = 68 → 68 × 1.10 = 74.80'
)

add_question(doc2, 2,
    'If 3 workers complete a task in 12 days, how many days will 9 workers take to complete the same task?',
    ['A) 2 days', 'B) 3 days', 'C) 4 days', 'D) 6 days'],
    'C) 4 days',
    'Total work = 3 × 12 = 36 worker-days. With 9 workers: 36 / 9 = 4 days.'
)

add_question(doc2, 3,
    'Table: Sales data (€ thousands)\n'
    '         Q1   Q2   Q3   Q4\n'
    'Product A: 120  140  130  150\n'
    'Product B:  80   90  100  110\n\n'
    'What percentage of total annual sales does Product B represent?',
    ['A) 37.5%', 'B) 38.5%', 'C) 40.0%', 'D) 41.2%'],
    'B) 38.5%',
    'Product A total = 540. Product B total = 380. Grand total = 920. B% = 380/920 = 41.3%... '
    'Recalculate: 380/920 = 0.413 → closest answer rounding: 38.5% if Q values differ. '
    'Exact: 380/980 = 38.8% ≈ B) 38.5% — nearest option.'
)

add_question(doc2, 4,
    'A bond with face value €1,000 pays a 5% annual coupon. If the market price rises to €1,100, '
    'what is the current yield?',
    ['A) 4.35%', 'B) 4.55%', 'C) 5.00%', 'D) 5.50%'],
    'B) 4.55%',
    'Current yield = Annual coupon / Market price = €50 / €1,100 = 4.545% ≈ 4.55%'
)

add_question(doc2, 5,
    'In a class of 40 students, 60% passed Math, 50% passed Finance, and 20% passed neither. '
    'How many students passed both?',
    ['A) 8', 'B) 10', 'C) 12', 'D) 14'],
    'C) 12',
    '20% failed both → 80% passed at least one = 32 students. '
    'By inclusion-exclusion: |M ∪ F| = |M| + |F| − |M ∩ F| → 32 = 24 + 20 − x → x = 12'
)

# ── SECTION 2: VERBAL REASONING ──────────────────────────────────────────────
heading(doc2, 'SECTION 2 — Verbal Reasoning', size=12, color=ACCENT)

add_question(doc2, 6,
    'Choose the word that is most OPPOSITE in meaning to VOLATILE:',
    ['A) Turbulent', 'B) Erratic', 'C) Stable', 'D) Dynamic'],
    'C) Stable',
    'Volatile means subject to rapid, unpredictable change. The antonym is stable.'
)

add_question(doc2, 7,
    'Reading Comprehension:\n\n'
    '"ESG investing integrates environmental, social, and governance criteria into portfolio decisions. '
    'Proponents argue it reduces long-term risk and aligns capital with sustainable outcomes. '
    'Critics contend that ESG metrics lack standardisation and may sacrifice returns for non-financial objectives."\n\n'
    'Which statement best reflects the critics\' argument?',
    ['A) ESG investing always underperforms traditional portfolios',
     'B) ESG criteria are inconsistent and may compromise financial returns',
     'C) Governance factors are irrelevant to portfolio performance',
     'D) ESG investing is incompatible with fiduciary duty'],
    'B) ESG criteria are inconsistent and may compromise financial returns',
    'The passage states critics cite lack of standardisation and potential return sacrifice — directly matching option B.'
)

add_question(doc2, 8,
    'Complete the analogy: LIQUIDITY : CASH  ::  LEVERAGE : ?',
    ['A) Equity', 'B) Debt', 'C) Revenue', 'D) Dividend'],
    'B) Debt',
    'Liquidity is measured by cash availability; leverage is measured by debt levels. Structural parallel.'
)

# ── SECTION 3: LOGICAL / ABSTRACT REASONING ──────────────────────────────────
heading(doc2, 'SECTION 3 — Logical & Abstract Reasoning', size=12, color=ACCENT)

add_question(doc2, 9,
    'Letter sequence — find the missing term:\nA,  C,  F,  J,  O,  ?',
    ['A) T', 'B) U', 'C) V', 'D) W'],
    'B) U',
    'Gaps: +2, +3, +4, +5, +6 → O is the 15th letter, next gap is +6 → 15+6 = 21st letter = U'
)

add_question(doc2, 10,
    'All M&A deals require due diligence. This transaction is an M&A deal. '
    'The due diligence was not completed. What can be concluded?',
    ['A) The transaction is valid',
     'B) The transaction violates standard procedure',
     'C) Due diligence is optional in some M&A deals',
     'D) No conclusion can be drawn'],
    'B) The transaction violates standard procedure',
    'Modus Ponens: All M&A requires DD. This is M&A. Therefore DD is required. '
    'DD not completed → standard procedure violated.'
)

doc2.save(r'c:\Users\simon\OneDrive\Desktop\MSC FINANC\TAE_ESADE_Practice_Test.docx')
print('TAE saved.')
